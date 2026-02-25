import os
import json
from typing import Dict, Any
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

DOCS_DIR = "generated_docs"
os.makedirs(DOCS_DIR, exist_ok=True)


def format_currency(amount) -> str:
    if not amount:
        return "не указано"
    return f"{amount:,} ₽".replace(",", " ")


async def generate_pdf(application_data: Dict[str, Any], grant_data: Dict[str, Any], output_path: str) -> str:
    """Generate PDF application document."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        font_path = os.path.join(os.path.dirname(__file__), "..", "fonts", "DejaVuSans.ttf")
        font_bold_path = os.path.join(os.path.dirname(__file__), "..", "fonts", "DejaVuSans-Bold.ttf")
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont("DejaVu", font_path))
            if os.path.exists(font_bold_path):
                pdfmetrics.registerFont(TTFont("DejaVu-Bold", font_bold_path))
            base_font = "DejaVu"
        else:
            base_font = "Helvetica"
        doc = SimpleDocTemplate(output_path, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("Title", fontName=base_font, fontSize=16, spaceAfter=12, alignment=1)
        heading_style = ParagraphStyle("Heading", fontName=base_font, fontSize=12, spaceAfter=6, spaceBefore=12, textColor=colors.darkblue)
        body_style = ParagraphStyle("Body", fontName=base_font, fontSize=10, spaceAfter=4)
        story = []
        wizard = application_data.get("wizard_data", {})
        story.append(Paragraph("ЗАЯВКА НА ГРАНТ", title_style))
        story.append(Paragraph(grant_data.get("source_name", ""), title_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.darkblue))
        story.append(Spacer(1, 0.5*cm))
        fmt = "%d.%m.%Y"
        story.append(Paragraph(f"Дата создания: {datetime.now().strftime(fmt)}", body_style))
        story.append(Spacer(1, 0.5*cm))
        applicant = wizard.get("step1", {})
        if applicant:
            story.append(Paragraph("1. СВЕДЕНИЯ О ЗАЯВИТЕЛЕ", heading_style))
            fields = [("ФИО", applicant.get("full_name", "")), ("Возраст", applicant.get("age", "")), ("Организация", applicant.get("organization", "")), ("ИНН", applicant.get("inn", "")), ("ОГРН", applicant.get("ogrn", "")), ("Регион", applicant.get("region", "")), ("Email", applicant.get("email", "")), ("Телефон", applicant.get("phone", ""))]
            for label, value in fields:
                if value:
                    story.append(Paragraph(f"<b>{label}:</b> {value}", body_style))
        project = wizard.get("step2", {})
        if project:
            story.append(Paragraph("2. ОПИСАНИЕ ПРОЕКТА", heading_style))
            for attr_lbl, attr_key in [("Название", "project_name"), ("Описание", "description"), ("Цели", "goals"), ("Ожидаемые результаты", "expected_results"), ("Целевая аудитория", "target_audience")]:
                val = project.get(attr_key, "")
                story.append(Paragraph(f"<b>{attr_lbl}:</b> {val}", body_style))
        budget = wizard.get("step3", {})
        if budget and budget.get("items"):
            story.append(Paragraph("3. БЮДЖЕТ ПРОЕКТА", heading_style))
            table_data = [["Категория", "Описание", "Сумма (₽)"]]
            total = 0
            for item in budget["items"]:
                amount = item.get("amount", 0)
                total += amount
                table_data.append([item.get("category", ""), item.get("description", ""), f"{amount:,}".replace(",", " ")])
            table_data.append(["ИТОГО", "", f"{total:,}".replace(",", " ")])
            table = Table(table_data, colWidths=[4*cm, 9*cm, 4*cm])
            table.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.darkblue), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("FONTNAME", (0,0), (-1,-1), base_font), ("FONTSIZE", (0,0), (-1,-1), 9), ("GRID", (0,0), (-1,-1), 0.5, colors.grey), ("BACKGROUND", (0,-1), (-1,-1), colors.lightgrey)]))
            story.append(table)
        doc.build(story)
        return output_path
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        raise


async def generate_docx(application_data: Dict[str, Any], grant_data: Dict[str, Any], output_path: str) -> str:
    """Generate DOCX application document."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    wizard = application_data.get("wizard_data", {})
    title = doc.add_heading("ЗАЯВКА НА ГРАНТ", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading(grant_data.get("source_name", ""), 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = "%d.%m.%Y"
    doc.add_paragraph(f"Дата: {datetime.now().strftime(fmt)}")
    doc.add_paragraph()
    applicant = wizard.get("step1", {})
    if applicant:
        doc.add_heading("1. СВЕДЕНИЯ О ЗАЯВИТЕЛЕ", 2)
        fields = [("ФИО", applicant.get("full_name", "")), ("Возраст", str(applicant.get("age", ""))), ("Организация", applicant.get("organization", "")), ("ИНН", applicant.get("inn", "")), ("ОГРН", applicant.get("ogrn", "")), ("Регион", applicant.get("region", "")), ("Email", applicant.get("email", "")), ("Телефон", applicant.get("phone", ""))]
        for label, value in fields:
            if value:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(value)
    project = wizard.get("step2", {})
    if project:
        doc.add_heading("2. ОПИСАНИЕ ПРОЕКТА", 2)
        for label, key in [("Название", "project_name"), ("Описание", "description"), ("Цели", "goals"), ("Ожидаемые результаты", "expected_results"), ("Целевая аудитория", "target_audience")]:
            if project.get(key):
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.bold = True
                p.add_run(project[key])
    budget = wizard.get("step3", {})
    if budget and budget.get("items"):
        doc.add_heading("3. БЮДЖЕТ ПРОЕКТА", 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Категория"
        hdr[1].text = "Описание"
        hdr[2].text = "Сумма (₽)"
        total = 0
        for item in budget["items"]:
            row = table.add_row().cells
            row[0].text = item.get("category", "")
            row[1].text = item.get("description", "")
            amount = item.get("amount", 0)
            total += amount
            row[2].text = f"{amount:,}".replace(",", " ")
        total_row = table.add_row().cells
        total_row[0].text = "ИТОГО"
        total_row[2].text = f"{total:,}".replace(",", " ")
    doc.save(output_path)
    return output_path
